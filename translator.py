import requests

from docx import Document
import os

subscription_key = "ENTER_YOUR_KEY"
endpoint = "ENTER YOUR ENDPOINT"
location = "ENTER YOUR LOCATION"
language_destination = 'pt-br'


def translator_text(text, target_language):
    path = '/translate'
    constructed_url = endpoint + path
    headers = {
        'Ocp-Apim-Subscription-Key': subscription_key,
        'Ocp-Apim-Subscription-Region': location,
        'Content-type': 'application/json',
        'X-ClientTraceId': str(os.urandom(16))
    }

    body = [{
        'text': text
    }]

    params = {
        'api-version': '3.0',
        'from':'en',
        'to':target_language,
    }
    request = requests.post(constructed_url, params=params, headers=headers, json=body)
    response = request.json()
    return response[0]["translations"][0]["text"]


def translate_document(path):
    document = Document(path)
    full_text = []
    for paragraph in document.paragraphs:
        translated_text = translator_text(paragraph.text, language_destination)
        full_text.append(translated_text)
    
    translated_doc = Document()
    for line in full_text:
        translated_doc.add_paragraph(line)
    path_translated = path.replace(".docx", f"_{language_destination}.docx")
    translated_doc.save(path)
    return path_translated